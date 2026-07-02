import os
from typing import List, Dict, Any
from pathlib import Path
from langchain_core.documents import Document
from utils.logger import logger
from .base_processor import BaseProcessor


class DOCXProcessor(BaseProcessor):
    """Process DOCX files and extract text"""

    def process(self, file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Process DOCX file
        
        Args:
            file_path: Path to DOCX file
            metadata: Optional metadata dictionary
            
        Returns:
            List of Document objects
        """
        documents = []
        metadata = metadata or {}
        
        try:
            from docx import Document as DocxDocument
            
            logger.info(f"Processing DOCX: {file_path}")
            
            doc = DocxDocument(file_path)
            content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            if content:
                full_text = "\n".join(content)
                doc_metadata = {
                    **metadata,
                    "source": Path(file_path).name,
                }
                
                document = Document(
                    page_content=full_text,
                    metadata=doc_metadata
                )
                documents.append(document)
            
            logger.info(f"Extracted {len(documents)} document(s) from DOCX")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
